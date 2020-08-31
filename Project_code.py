# -*- coding: utf-8 -*-
"""
Created on Tue Apr 28 23:15:00 2020

@author: Anurag
"""
### importing libraries

import os
import textract
from Crypto.Cipher import AES
from Crypto.Hash import SHA256, HMAC
import hashlib
import json
from base64 import b64encode,b64decode
from Crypto.Util.Padding import pad
import docx

### describe path of working directory
os.chdir("D:\masters data science\Spring session\Big Data Security (ISQS-5342-001)\Project")
my_data = textract.process("R11659082_data.docx")
### for SHA256
hash_object = SHA256.new(data=b'my_data')
hash_object_SHA=hash_object.hexdigest()
#print(hash_object.hexdigest())
### for HMAC
hash_object = HMAC.new(b'my_data')
hash_object_HMAC=hash_object.hexdigest()
#print(hash_object.hexdigest())



### AES
###encryption
key = b'1234567899876543'
cipher = AES.new(key, AES.MODE_CBC)
ct_bytes = cipher.encrypt(pad(my_data, AES.block_size))
iv = b64encode(cipher.iv).decode('utf-8')
ct = b64encode(ct_bytes).decode('utf-8')
result = json.dumps({'iv':iv, 'ciphertext':ct})
print(result)

## exporting results to a bin file
with open('R11659082_AES.bin', 'wb') as f:
    f.write(result.encode())

####decryption

with open("R11659082_AES.bin","rb") as f:
    json_input= f.read()

#json_input=result
from Crypto.Util.Padding import unpad

b64 = json.loads(json_input)
iv = b64decode(b64['iv'])
ct = b64decode(b64['ciphertext'])
cipher = AES.new(key, AES.MODE_CBC, iv)
pt = unpad(cipher.decrypt(ct), AES.block_size)
print("The message was: ", pt)


temp = docx.Document()
temp.add_paragraph(pt.decode("utf-8"))
temp.save("R11659082_AES_decrypted.docx")

    
######################## Salsa   
###encryption
from Crypto.Cipher import Salsa20

cipher = Salsa20.new(key)
msg = cipher.nonce + cipher.encrypt(my_data)    
print(msg)
with open('R11659082_Salsa.bin', 'wb') as f:
    f.write(msg)

### decryption
with open('R11659082_Salsa.bin', 'rb') as f:
    msg= f.read()
msg_nonce = msg[:8]
ciphertext = msg[8:]
cipher = Salsa20.new(key, nonce=msg_nonce)
plaintext = cipher.decrypt(ciphertext)
print(plaintext)
temp = docx.Document()
temp.add_paragraph(plaintext.decode("utf-8"))
temp.save("R11659082_Salsa_decrypted.docx")
 
######## for comparison 
my_data = textract.process("R11659082_AES_decrypted.docx")    
### for SHA256
hash_object_decrypted = SHA256.new(data=b'my_data')
hash_object_decrypted= hash_object_decrypted.hexdigest()

if (hash_object_SHA == hash_object_decrypted):
    print ("Hurray! matched")
else :
    print("Comparison Failed")    
    
     
### for HMAC
my_data = textract.process("R11659082_AES_decrypted.docx")    
    
hash_object_d = HMAC.new(b'my_data')
hash_object_d=hash_object_d.hexdigest() 
if (hash_object_HMAC == hash_object_d):
    print ("Hurray! matched")
else :
    print("Comparison Failed")
    
    
 #### for comparison using salsa20   
    
my_data = textract.process("R11659082_Salsa_decrypted.docx")    
### for SHA256
hash_object_decrypted = SHA256.new(data=b'my_data')
hash_object_decrypted= hash_object_decrypted.hexdigest()

if (hash_object_SHA == hash_object_decrypted):
    print ("Hurray! matched")
else :
    print("Comparison Failed")    
    
     
### for HMAC
my_data = textract.process("R11659082_Salsa_decrypted.docx")    
    
hash_object_d = HMAC.new(b'my_data')
hash_object_d=hash_object_d.hexdigest() 
if (hash_object_HMAC == hash_object_d):
    print ("Hurray! matched")
else :
    print("Comparison Failed")